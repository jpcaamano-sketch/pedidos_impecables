import google.generativeai as genai
from docx import Document
import io

# --- 1. CONFIGURACIÓN ---
st.set_page_config(page_title="Generador de Pedidos Impecables", page_icon="🗣️", layout="centered")

# CSS para ocultar encabezados, pie de página y menú, y estilizar la app
st.markdown("""
    <style>
    /* Estilos para inputs más grandes */
    .stTextArea textarea { font-size: 16px !important; }
    .stTextInput input { font-size: 16px !important; }
    .css-1v0mbdj { width: 100%; }
    .info-box { background-color: #f0f8ff; padding: 15px; border-radius: 10px; border-left: 5px solid #1f77b4; }
    
    /* --- OCULTAR ELEMENTOS DE LA INTERFAZ DE STREAMLIT --- */
    #MainMenu {visibility: hidden;} /* Oculta el menú de hamburguesa (derecha arriba) */
    header {visibility: hidden;}    /* Oculta la barra de cabecera superior */
    footer {visibility: hidden;}    /* Oculta el pie de página "Made with Streamlit" */
    </style>
    """, unsafe_allow_html=True)

# --- 2. CONEXIÓN IA ---
try:
    api_key = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=api_key)
except Exception:
    st.error("⚠️ Falta la API KEY en .streamlit/secrets.toml")
    st.stop()

# --- 3. LÓGICA DE PEDIDOS (ONTOLOGÍA DEL LENGUAJE) ---
def generar_pedido(oyente, accion, condiciones, tiempo, contexto):
    try:
        model = genai.GenerativeModel("gemma-3-27b-it")
        
        prompt = f"""
        Actúa como un Coach Ontológico experto en Fernando Flores y Rafael Echeverría.
        Tu tarea es redactar un "PEDIDO IMPECABLE" (Speech Act) basado en estos datos.

        DATOS DEL PEDIDO:
        1. Oyente: {oyente}
        2. Acción futura: {accion}
        3. Condiciones de Satisfacción (Estándar de calidad): {condiciones}
        4. Factor Tiempo: {tiempo}
        5. Trasfondo (Por qué es importante): {contexto}

        ESTRUCTURA DE RESPUESTA:
        Genera dos secciones:

        SECCION_GUION:
        Escribe el guion conversacional exacto, en primera persona, listo para ser hablado o enviado.
        El tono debe ser asertivo pero colaborativo. 
        IMPORTANTE: Debe terminar explícitamente buscando la aceptación del otro (Ej: "¿Puedes comprometerte a esto?", "¿Cuento contigo?").
        
        SECCION_ANALISIS:
        Explica brevemente por qué este pedido reduce la incertidumbre, destacando cómo las condiciones de satisfacción evitan malentendidos.
        """
        
        response = model.generate_content(prompt)
        parts = response.text.split("SECCION_ANALISIS:")
        
        guion = parts[0].replace("SECCION_GUION:", "").strip()
        analisis = parts[1].strip() if len(parts) > 1 else "Análisis no generado."
        
        return guion, analisis

    except Exception as e:
        return f"Error: {e}", ""

def crear_docx(guion, analisis):
    doc = Document()
    doc.add_heading('Guion de Pedido Impecable', 0)
    
    doc.add_heading('Conversación Sugerida:', level=1)
    doc.add_paragraph(guion)
    
    doc.add_heading('Análisis Ontológico:', level=1)
    doc.add_paragraph(analisis)
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- 4. INTERFAZ ---
st.title("🗣️ Pedidos Impecables")
st.markdown("**Basado en la Ontología del Lenguaje (Fernando Flores)**")
st.caption("Un pedido no es un deseo. Es una acción lingüística que coordina acciones futuras.")

st.divider()

with st.container(border=True):
    st.subheader("🛠️ Diseña tu Pedido")
    
    col1, col2 = st.columns(2)
    with col1:
        oyente = st.text_input("1. ¿A quién le pides? (Oyente)", placeholder="Ej: Juan, Jefe de Proyecto")
    with col2:
        tiempo = st.text_input("2. Factor Tiempo (¿Para cuándo?)", placeholder="Ej: Martes 15 a las 14:00 hrs")

    accion = st.text_area("3. Acción (¿Qué quieres que haga?)", placeholder="Ej: Que prepares el reporte de ventas...")
    
    st.info("💡 **Clave del Éxito:** Las condiciones de satisfacción eliminan la frase 'es que yo pensé que...'")
    condiciones = st.text_area("4. Condiciones de Satisfacción (¿Cómo sabes que está bien hecho?)", 
                               placeholder="Ej: Debe estar en formato Excel, incluir el IVA desglosado y tener máximo 2 páginas.", height=100)
    
    contexto = st.text_area("5. Trasfondo (¿Cuál es el quiebre/necesidad?)", 
                            placeholder="Ej: Tenemos reunión de directorio el miércoles y necesito datos duros para defender el presupuesto.")

    if st.button("✨ Generar Pedido Impecable", type="primary"):
        if not oyente or not accion or not condiciones or not tiempo:
            st.warning("⚠️ Para que el pedido sea impecable, necesitas llenar todos los campos (especialmente las condiciones y el tiempo).")
        else:
            with st.spinner("Construyendo acto del habla..."):
                guion_gen, analisis_gen = generar_pedido(oyente, accion, condiciones, tiempo, contexto)
                st.session_state.pedido = {"guion": guion_gen, "analisis": analisis_gen}

# --- 5. RESULTADOS ---
if 'pedido' in st.session_state:
    res = st.session_state.pedido
    
    st.divider()
    st.subheader("💬 Tu Guion")
    
    st.markdown(f"""
    <div class="info-box">
        {res['guion'].replace(chr(10), '<br>')}
    </div>
    """, unsafe_allow_html=True)
    
    with st.expander("🧠 Ver Análisis Ontológico (Por qué funciona)"):
        st.write(res['analisis'])
    
    # Descarga
    docx_file = crear_docx(res['guion'], res['analisis'])
    st.download_button(
        label="💾 Descargar Guion (.docx)",
        data=docx_file,
        file_name="Pedido_Impecable.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
